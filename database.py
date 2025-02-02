import sqlite3
import random
import requests
import os
import pandas as pd
from datetime import datetime
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings

warnings.filterwarnings("ignore")

USAJOBS_API_KEY = "HiVRyd6XtR3ynF1tWr3q4qPQnJ63xF1dPPvIEjebkhc="
USAJOBS_API_URL = "https://data.usajobs.gov/api/Search"

ADZUNA_API_KEY = "54a9e1822eb536ed93013900a73a68c3"
ADZUNA_API_URL = "https://api.adzuna.com/v1/api/jobs/us/search/1"
ADZUNA_API_ID = "9d1bc7b6"

# ---------------- API CONFIGURATION ----------------
def get_usajobs_config():
    """ Fetches USAJobs API key and URL from environment variables """
    return {
        "api_key": "HiVRyd6XtR3ynF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1tWr3q4qPQnJ63xF1t",
        "api_url": "https://data.usajobs.gov/api/Search",
    }

def get_adzuna_config():
    """ Fetches Adzuna API credentials from environment variables """
    return {
        "app_id": "9d1bc7b6",
        "api_key": "54a9e1822eb536ed93013900a73a68c3",
        "api_url": "https://api.adzuna.com/v1/api/jobs/us/search/1",
    }

# ---------------- USER-DEFINED SEARCH PARAMETERS ----------------
def get_search_parameters():
    """ Fetches job search parameters from user input (provided via app.py) """
    return {
        "keywords": [],
        "skills": [],
        "location": "United States",
        "results_per_page": 25,
    }

# ---------------- DATABASE CONNECTION ----------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "job_recommendations.db")

def create_connection():
    """Establishes a connection to the SQLite database inside the Hackathon folder."""
    return sqlite3.connect(DB_PATH, timeout=5)  # Prevents database lock errors

# ---------------- CREATE TABLES ----------------
def create_tables():
    """Creates all necessary tables in the database if they do not exist."""
    conn = create_connection()
    cursor = conn.cursor()
    
    # Enable Write-Ahead Logging (WAL) mode
    cursor.execute("PRAGMA journal_mode=WAL;")

    tables = {
        "USER": """
            CREATE TABLE IF NOT EXISTS USER (
                USER_ID TEXT PRIMARY KEY, 
                USER_USERNAME TEXT UNIQUE NOT NULL,
                USER_RESUME TEXT
            );
        """,
        "JOB": """
            CREATE TABLE IF NOT EXISTS JOB (
                JOB_ID TEXT PRIMARY KEY, 
                JOB_TITLE TEXT NOT NULL, 
                JOB_COMPANY TEXT NOT NULL, 
                JOB_LOCATION TEXT NOT NULL, 
                JOB_SALARY FLOAT(10,2), 
                JOB_DATEPOSTED TEXT NOT NULL, 
                JOB_DESCRIPTION TEXT, 
                JOB_CONTACTEMAIL TEXT, 
                JOB_REQUIREMENT TEXT, 
                JOB_URL TEXT NOT NULL, 
                JOB_API TEXT NOT NULL, 
                JOB_RECOMMENDED INTEGER DEFAULT 0, 
                JOB_REVIEWCOUNT INTEGER DEFAULT 0, 
                JOB_REVIEWRATING FLOAT DEFAULT 0.0
            );
        """,
        "KEYWORD": """
            CREATE TABLE IF NOT EXISTS KEYWORD (
                KEYWORD_ID TEXT PRIMARY KEY, 
                KEYWORD_NAME TEXT UNIQUE NOT NULL
            );
        """,
        "SKILL": """
            CREATE TABLE IF NOT EXISTS SKILL (
                SKILL_ID TEXT PRIMARY KEY, 
                SKILL_NAME TEXT UNIQUE NOT NULL
            );
        """,
        "USER_KEYWORD": """
            CREATE TABLE IF NOT EXISTS USER_KEYWORD (
                USER_ID TEXT, 
                KEYWORD_ID TEXT, 
                PRIMARY KEY (USER_ID, KEYWORD_ID),
                FOREIGN KEY (USER_ID) REFERENCES USER (USER_ID),
                FOREIGN KEY (KEYWORD_ID) REFERENCES KEYWORD (KEYWORD_ID)
            );
        """,
        "USER_SKILL": """
            CREATE TABLE IF NOT EXISTS USER_SKILL (
                USER_ID TEXT, 
                SKILL_ID TEXT, 
                PRIMARY KEY (USER_ID, SKILL_ID),
                FOREIGN KEY (USER_ID) REFERENCES USER (USER_ID),
                FOREIGN KEY (SKILL_ID) REFERENCES SKILL (SKILL_ID)
            );
        """,
        "JOB_SKILL": """
            CREATE TABLE IF NOT EXISTS JOB_SKILL (
                JOB_ID TEXT, 
                SKILL_ID TEXT, 
                PRIMARY KEY (JOB_ID, SKILL_ID),
                FOREIGN KEY (JOB_ID) REFERENCES JOB (JOB_ID),
                FOREIGN KEY (SKILL_ID) REFERENCES SKILL (SKILL_ID)
            );
        """,
        "RECOMMEND": """
            CREATE TABLE IF NOT EXISTS RECOMMEND (
                USER_ID TEXT,
                JOB_ID TEXT,
                REC_DATE TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                REC_RATED INTEGER DEFAULT 0,  -- 0 = Not Rated, 1 = Rated
                REC_RATING INTEGER CHECK (REC_RATING BETWEEN 1 AND 5) DEFAULT NULL,  -- User rating (1-5 or NULL)
                PRIMARY KEY (USER_ID, JOB_ID, REC_DATE),
                FOREIGN KEY (USER_ID) REFERENCES USER (USER_ID),
                FOREIGN KEY (JOB_ID) REFERENCES JOB (JOB_ID)
            );
        """
    }

    for table_name, query in tables.items():
        cursor.execute(query)

    conn.commit()
    conn.close()

# ---------------- GENERATE UNIQUE IDs ----------------
def generate_id(prefix):
    """Generates a unique ID with a given prefix (e.g., U12345, S67890)."""
    return f"{prefix}{random.randint(10000, 99999)}"

def generate_job_id(job_title, job_company, job_dateposted):
    """Generates a unique JOB_ID based on job title, company, and date posted."""
    formatted_title = job_title.replace(" ", "").lower()
    formatted_company = job_company.replace(" ", "").lower()

    try:
        formatted_date = datetime.strptime(job_dateposted[:10], "%Y-%m-%d").strftime("%d-%m-%y")
    except ValueError:
        formatted_date = "unknown"

    return f"{formatted_title}{formatted_company}{formatted_date}"

# ---------------- INSERT INTO TABLES ----------------
def insert_user(username, resume_text=None):
    """Inserts a new user with a unique username and optional resume text."""
    conn = create_connection()
    cursor = conn.cursor()

    # Check if username already exists
    cursor.execute("SELECT USER_ID FROM USER WHERE USER_USERNAME = ?", (username,))
    existing_user = cursor.fetchone()

    if existing_user:
        user_id = existing_user[0]
    else:
        user_id = generate_id("U")
        cursor.execute("INSERT INTO USER (USER_ID, USER_USERNAME, USER_RESUME) VALUES (?, ?, ?)",
                       (user_id, username, resume_text))
        conn.commit()

    conn.close()
    return user_id  # Ensure the correct user ID is returned

def insert_skills(skills):
    """Inserts a list of skills into the SKILL table, ignoring duplicates."""
    conn = create_connection()
    cursor = conn.cursor()

    for skill in skills:
        skill_id = generate_id("S")
        cursor.execute("INSERT OR IGNORE INTO SKILL (SKILL_ID, SKILL_NAME) VALUES (?, ?)", (skill_id, skill))

    conn.commit()
    conn.close()
    
def assign_user_skills(user_id, user_skills):
    """Assigns skills to a user based on their input."""
    if not user_skills:
        return  # Prevent empty inserts

    conn = create_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT SKILL_ID, SKILL_NAME FROM SKILL")
    skill_map = {name: sid for sid, name in cursor.fetchall()}  # Map skills to IDs

    user_skill_pairs = [(user_id, skill_map[skill]) for skill in user_skills if skill in skill_map]
    
    if user_skill_pairs:
        cursor.executemany("INSERT OR IGNORE INTO USER_SKILL (USER_ID, SKILL_ID) VALUES (?, ?)", user_skill_pairs)
        conn.commit()  # Ensure data is stored

    conn.close()

def insert_keyword(keyword):
    """ Inserts a keyword if it doesn't already exist and returns its ID. """
    conn = create_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("SELECT KEYWORD_ID FROM KEYWORD WHERE KEYWORD_NAME = ?", (keyword,))
        existing = cursor.fetchone()

        if existing:
            return existing[0]  # Return existing keyword ID

        keyword_id = generate_id("K")
        cursor.execute("INSERT INTO KEYWORD (KEYWORD_ID, KEYWORD_NAME) VALUES (?, ?)", (keyword_id, keyword))
        conn.commit()
        return keyword_id  # Return new keyword ID

    except sqlite3.Error as e:
        print(f"SQLite Error in insert_keyword: {e}")
        return None

    finally:
        cursor.close()
        conn.close()

def insert_user_keywords(user_id, keyword_ids):
    """ Inserts multiple keyword relationships for a user in one batch. """
    if not keyword_ids:
        return

    conn = create_connection()
    cursor = conn.cursor()
    
    try:
        user_keyword_pairs = [(user_id, keyword_id) for keyword_id in keyword_ids]
        cursor.executemany(
            "INSERT OR IGNORE INTO USER_KEYWORD (USER_ID, KEYWORD_ID) VALUES (?, ?)", user_keyword_pairs
        )
        conn.commit()
    except sqlite3.Error as e:
        print(f"SQLite Error in insert_user_keywords: {e}")
    finally:
        cursor.close()
        conn.close()
            
# ---------------- JOB RECOMMENDATION STATS ----------------
def update_job_recommendation_stats(job_id):
    """ Updates JOB_RECOMMENDED, JOB_REVIEWCOUNT, and JOB_REVIEWRATING for a given job. """
    conn = create_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT COUNT(*) FROM RECOMMEND WHERE JOB_ID = ?", (job_id,))
    total_recommended = cursor.fetchone()[0] or 0

    cursor.execute("SELECT SUM(REC_RATED) FROM RECOMMEND WHERE JOB_ID = ?", (job_id,))
    total_review_count = cursor.fetchone()[0] or 0

    cursor.execute("SELECT AVG(REC_RATING) FROM RECOMMEND WHERE JOB_ID = ? AND REC_RATING IS NOT NULL", (job_id,))
    avg_review_rating = cursor.fetchone()[0]

    avg_review_rating = round(avg_review_rating, 2) if avg_review_rating is not None else 0.0

    cursor.execute("""
        UPDATE JOB 
        SET JOB_RECOMMENDED = ?, 
            JOB_REVIEWCOUNT = ?, 
            JOB_REVIEWRATING = ?
        WHERE JOB_ID = ?
    """, (total_recommended, total_review_count, avg_review_rating, job_id))

    conn.commit()
    conn.close()

# ---------------- INSERT RECOMMENDATION ----------------
def insert_recommendation(user_id, job_id):
    """ Inserts a job recommendation for a user and updates job stats. """
    conn = create_connection()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            INSERT INTO RECOMMEND (USER_ID, JOB_ID) 
            VALUES (?, ?)
        """, (user_id, job_id))
        
        conn.commit()
        update_job_recommendation_stats(job_id)  
    except sqlite3.IntegrityError:
        print(f"⚠️ Job {job_id} was already recommended to user {user_id}.")

    conn.close()

# ---------------- INSERT JOB SKILLS ----------------
def insert_job_skills(job_id, job_skills):
    """ Inserts multiple skill relationships for a job in one batch. """
    if not job_skills:
        return

    conn = create_connection()
    cursor = conn.cursor()
    
    try:
        # Get SKILL_IDs from the SKILL table
        cursor.execute(
            "SELECT SKILL_ID FROM SKILL WHERE SKILL_NAME IN ({})".format(",".join("?" * len(job_skills))),
            job_skills
        )
        skill_ids = [row[0] for row in cursor.fetchall()]

        # Create job-skill pairs
        job_skill_pairs = [(job_id, skill_id) for skill_id in skill_ids]

        # Batch insert into JOB_SKILL
        cursor.executemany(
            "INSERT OR IGNORE INTO JOB_SKILL (JOB_ID, SKILL_ID) VALUES (?, ?)", job_skill_pairs
        )
        conn.commit()
    except sqlite3.Error as e:
        print(f"SQLite Error in insert_job_skills: {e}")
    finally:
        cursor.close()
        conn.close()

# ---------------- INSERT JOB DATA ----------------
def insert_job_data(job_list):
    """ Inserts job listings into the JOB table and assigns skills. """
    conn = create_connection()
    cursor = conn.cursor()

    for job in job_list:
        cursor.execute("SELECT JOB_ID FROM JOB WHERE JOB_ID = ?", (job["id"],))
        if cursor.fetchone() is None:
            cursor.execute("""
                INSERT INTO JOB (JOB_ID, JOB_TITLE, JOB_COMPANY, JOB_LOCATION, JOB_SALARY, JOB_DATEPOSTED,
                                 JOB_DESCRIPTION, JOB_CONTACTEMAIL, JOB_REQUIREMENT, JOB_URL, JOB_API,
                                 JOB_RECOMMENDED, JOB_REVIEWCOUNT, JOB_REVIEWRATING)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (job["id"], job["title"], job["company"], job["location"], job["salary"],
                  job["date_posted"], job["description"], job["contact_email"],
                  job["requirements"], job["url"], job["api_source"], 0, 0, 0.0))

            conn.commit()  # Ensure job data is committed before adding skills

            # Assign skills to the job if provided
            if "skills" in job and job["skills"]:
                insert_job_skills(job["id"], job["skills"])

    conn.close()

# ---------------- RATE RECOMMENDATION ----------------
def rate_recommendation(user_id, job_id, rating):
    """ Updates a recommended job with a user rating (1-5) and updates job stats. """
    if rating < 1 or rating > 5:
        print("⚠️ Invalid rating. Please enter a value between 1 and 5.")
        return

    conn = create_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM RECOMMEND WHERE USER_ID = ? AND JOB_ID = ?", (user_id, job_id))
    if cursor.fetchone():
        cursor.execute("""
            UPDATE RECOMMEND 
            SET REC_RATED = 1, REC_RATING = ? 
            WHERE USER_ID = ? AND JOB_ID = ?
        """, (rating, user_id, job_id))

        conn.commit()
        update_job_recommendation_stats(job_id)  
        print(f"✅ Rating {rating} saved for job {job_id}!")
    else:
        print(f"⚠️ Job {job_id} was not recommended to user {user_id}.")

    conn.close()
    
# ---------------- FETCH USER RESUME ----------------
def get_user_resume(user_id):
    """ Fetches user resume text from the database. """
    conn = create_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT USER_RESUME FROM USER WHERE USER_ID = ?", (user_id,))
    resume = cursor.fetchone()
    conn.close()
    return resume[0] if resume else None
    
# ---------------- FETCH JOB LISTINGS ----------------
def fetch_usajobs(keyword, location, results_per_page=25):
    """ Fetch jobs from USAJobs API with error handling based on user input. """
    headers = {
        "Host": "data.usajobs.gov",
        "User-Agent": "YourEmail@example.com",
        "Authorization-Key": USAJOBS_API_KEY
    }
    params = {"Keyword": keyword, "LocationName": location, "ResultsPerPage": results_per_page}

    response = requests.get(USAJOBS_API_URL, headers=headers, params=params)
    
    if response.status_code == 200:
        try:
            jobs = response.json().get("SearchResult", {}).get("SearchResultItems", [])
            job_list = []
            for job in jobs:
                job_data = job.get("MatchedObjectDescriptor", {})

                description = job_data.get("UserArea", {}).get("Details", {}).get("MajorDuties", "N/A")
                if isinstance(description, list):
                    description = "\n".join(description)

                remuneration = job_data.get("PositionRemuneration", [{}])
                
                job_list.append({
                    "id": generate_job_id(job_data.get("PositionTitle", ""), job_data.get("OrganizationName", ""), job_data.get("PublicationStartDate", "")),
                    "title": job_data.get("PositionTitle", "N/A"),
                    "company": job_data.get("OrganizationName", "N/A"),
                    "location": job_data.get("PositionLocationDisplay", location),
                    "salary": remuneration[0].get("MinimumRange", "N/A") if remuneration else "N/A",
                    "date_posted": job_data.get("PublicationStartDate", "N/A"),
                    "description": description,
                    "contact_email": job_data.get("UserArea", {}).get("Details", {}).get("AgencyContactEmail", "N/A"),
                    "requirements": job_data.get("QualificationSummary", "N/A"),
                    "url": job_data.get("PositionURI", "N/A"),
                    "api_source": "USAJobs"
                })
            return job_list
        except Exception as e:
            print("Error parsing USAJobs response:", e)
            return []
    else:
        print(f"USAJobs API Error {response.status_code}: {response.text}")
        return []

def fetch_adzuna_jobs(keyword, location, results_per_page=25):
    """ Fetch jobs from Adzuna API with user-defined search parameters. """
    params = {
        "app_id": ADZUNA_API_ID, "app_key": ADZUNA_API_KEY,
        "results_per_page": results_per_page, "what": keyword, "where": location
    }

    response = requests.get(ADZUNA_API_URL, params=params)
    
    if response.status_code == 200:
        try:
            jobs = response.json().get("results", [])
            job_list = []
            for job in jobs:
                job_list.append({
                    "id": generate_job_id(job.get("title", ""), job.get("company", {}).get("display_name", "N/A"), job.get("created", "")),
                    "title": job.get("title", "N/A"),
                    "company": job.get("company", {}).get("display_name", "N/A"),
                    "location": job.get("location", {}).get("display_name", location),
                    "salary": job.get("salary_min", "N/A"),
                    "date_posted": job.get("created", "N/A"),
                    "description": job.get("description", "N/A"),
                    "contact_email": "N/A",
                    "requirements": "N/A",
                    "url": job.get("redirect_url", "N/A"),
                    "api_source": "Adzuna"
                })
            return job_list
        except Exception as e:
            print("Error parsing Adzuna response:", e)
            return []
    else:
        print(f"Adzuna API Error {response.status_code}: {response.text}")
        return []

# ---------------- COMBINE JOB LISTINGS ----------------
def get_combined_jobs(keywords, location='United States', results_per_page=25):
    """ Fetches jobs from both APIs based on user input. """
    all_jobs = []
    for keyword in keywords:
        all_jobs.extend(fetch_usajobs(keyword, location, results_per_page))
        all_jobs.extend(fetch_adzuna_jobs(keyword, location, results_per_page))
    return all_jobs

# ---------------- READ RESUME FROM .DOCX ----------------
def read_resume(file_path):
    """ Reads text from a .docx resume file provided by the user. """
    if not os.path.exists(file_path):
        print(f"⚠️ Resume file not found: {file_path}")
        return None

    try:
        doc = Document(file_path)
        return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        print(f"Error reading resume file: {e}")
        return None

# ---------------- STORE RESUME IN DATABASE ----------------
def store_resume(user_id, file_path):
    """ Stores the resume content into the database from a user-uploaded file. """
    if not os.path.exists(file_path):
        return {"error": f"Resume file not found at {file_path}"}

    try:
        doc = Document(file_path)
        resume_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        return {"error": f"Failed to read resume file: {str(e)}"}

    if not resume_text:
        return {"error": "Extracted resume text is empty. Ensure the document is not blank."}

    try:
        conn = sqlite3.connect("job_recommendations.db")
        cursor = conn.cursor()
        cursor.execute("UPDATE USER SET USER_RESUME = ? WHERE USER_ID = ?", (resume_text, user_id))
        conn.commit()
        conn.close()
        return {"success": "Resume stored successfully"}
    except sqlite3.Error as e:
        return {"error": f"Database error: {str(e)}"}

# ---------------- AI MODEL: RECOMMEND JOBS ----------------
def get_jobs():
    """ Fetch job listings from the database """
    conn = create_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT JOB_ID, JOB_TITLE, JOB_DESCRIPTION FROM JOB")
    jobs = cursor.fetchall()
    conn.close()
    return jobs if jobs else []

def get_seen_jobs(user_id):
    """ Fetches JOB_IDs that have already been recommended to the user """
    conn = create_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT JOB_ID FROM RECOMMEND WHERE USER_ID = ?", (user_id,))
    seen_jobs = {row[0] for row in cursor.fetchall()}
    conn.close()
    return seen_jobs

def recommend_jobs(user_id, top_n=5, randomize=False):
    """ Uses AI to recommend jobs based on resume similarity, ensuring only new top N jobs are recommended.
        If randomize=True, it will randomly select from a larger pool of top recommendations.
    """
    resume_text = get_user_resume(user_id)
    if not resume_text:
        print("⚠️ No resume found for this user. Cannot generate recommendations.")
        return []

    job_data = get_jobs()
    if not job_data:
        print("⚠️ No jobs available for recommendation.")
        return []

    job_ids, job_titles, job_descriptions = zip(*job_data)
    # Ensure job descriptions are valid strings
    job_descriptions = [desc if isinstance(desc, str) else "" for desc in job_descriptions]

    # Compute TF-IDF vectors and cosine similarity scores
    vectorizer = TfidfVectorizer(stop_words="english")
    tfidf_matrix = vectorizer.fit_transform([resume_text] + list(job_descriptions))
    similarity_scores = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()

    # Rank jobs by similarity (highest first)
    ranked_jobs = sorted(zip(job_ids, job_titles, similarity_scores), key=lambda x: x[2], reverse=True)

    # Filter out jobs that have already been recommended
    seen_jobs = get_seen_jobs(user_id)
    new_recommendations = [(job_id, title, score) for job_id, title, score in ranked_jobs if job_id not in seen_jobs]

    # If randomize flag is true, randomly select from the top pool (say, top 20)
    if randomize:
        pool_size = min(20, len(new_recommendations))
        top_pool = new_recommendations[:pool_size]
        import random  # ensure random is imported
        random.shuffle(top_pool)
        recommended_jobs = top_pool[:top_n]
    else:
        recommended_jobs = new_recommendations[:top_n]

    if recommended_jobs:
        print("\n🔥 Top Recommended Jobs 🔥")
        for job_id, title, score in recommended_jobs:
            print(f"📌 {title} (Score: {score:.4f})")
            # Store this recommendation in the database so it won't show again in this session
            insert_recommendation(user_id, job_id)
    else:
        print("⚠️ No new job recommendations available.")

    return recommended_jobs

def regenerate_recommendations(user_id, top_n=5):
    """ Allows the user to get new recommendations different from previous ones """
    print("\n🔄 Regenerating Recommendations...")
    conn = create_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM RECOMMEND WHERE USER_ID = ?", (user_id,))
    conn.commit()
    conn.close()
    # Call recommend_jobs with randomize=True so that new selections are made
    return recommend_jobs(user_id, top_n, randomize=True)

# ✅ Load BLS Occupation Data
BLS_CSV_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bls_occupations_profiles.csv")
bls_data = pd.read_csv(BLS_CSV_PATH)
bls_data = bls_data.fillna("-")  # Replace NaN with empty string

def recommend_career_goal(user_id):
    """Analyzes user's resume, skills, and keywords to suggest a career goal from BLS data."""
    
    # Step 1: Fetch user resume
    conn = create_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT USER_RESUME FROM USER WHERE USER_ID = ?", (user_id,))
    resume_text = cursor.fetchone()
    resume_text = resume_text[0] if resume_text else ""

    # Step 2: Fetch user skills
    cursor.execute("""
        SELECT SKILL.SKILL_NAME FROM SKILL
        JOIN USER_SKILL ON SKILL.SKILL_ID = USER_SKILL.SKILL_ID
        WHERE USER_SKILL.USER_ID = ?
    """, (user_id,))
    user_skills = [row[0] for row in cursor.fetchall()]

    # Step 3: Fetch user keywords
    cursor.execute("""
        SELECT KEYWORD.KEYWORD_NAME FROM KEYWORD
        JOIN USER_KEYWORD ON KEYWORD.KEYWORD_ID = USER_KEYWORD.KEYWORD_ID
        WHERE USER_KEYWORD.USER_ID = ?
    """, (user_id,))
    user_keywords = [row[0] for row in cursor.fetchall()]

    conn.close()  # ✅ Close database connection

    # Ensure there's enough data to analyze
    if not resume_text and not user_skills and not user_keywords:
        return {"error": "Insufficient data to generate a career goal."}

    # Step 4: Process BLS Data
    bls_data["combined_info"] = (
        bls_data["OCC_TITLE"] + " " + 
        bls_data["DUTIES"] + " " +
        bls_data["EDUCATION"] + " " + 
        bls_data["EXPERIENCE"] + " " +
        bls_data["MEDIAN_PAY_YEAR"] + " " +
        bls_data["JOB_OUTLOOK"]
    )

    # Step 5: Compute Similarity
    vectorizer = TfidfVectorizer(stop_words="english")
    corpus = [" ".join([resume_text] + user_skills + user_keywords)] + list(bls_data["combined_info"])
    tfidf_matrix = vectorizer.fit_transform(corpus)
    
    # Step 6: Rank Career Matches
    similarity_scores = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
    bls_data["similarity"] = similarity_scores
    best_match = bls_data.sort_values(by="similarity", ascending=False).iloc[0]

    # ✅ Fix: Ensure correct column names when selecting best match
    career_goal = {
        "title": best_match["OCC_TITLE"],  # ✅ Fix: Replacing "title" with actual column name
        "duties": best_match["DUTIES"],
        "average_salary": best_match["MEDIAN_PAY_YEAR"],
        "education_needed": best_match["EDUCATION"],
        "experience_needed": best_match["EXPERIENCE"],
        "job_outlook": best_match["JOB_OUTLOOK"],
        "roadmap": f"To become a {best_match['OCC_TITLE']}, you need {best_match['EDUCATION']} "
                f"and {best_match['EXPERIENCE']} experience. Consider getting certifications in relevant fields. "
                f"The job outlook is {best_match['JOB_OUTLOOK']}, with an average salary of {best_match['MEDIAN_PAY_YEAR']}."
    }

    return career_goal

def format_salary(salary):
    """Formats a salary with $ sign, comma separator, and 2 decimal places."""
    try:
        salary = float(salary)
        return f"${salary:,.2f}"
    except (ValueError, TypeError):
        return "N/A"

if __name__ == "__main__":
    create_tables()
